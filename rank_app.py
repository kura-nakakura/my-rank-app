import streamlit as st
import streamlit.components.v1 as components
from google import genai
import re
from pypdf import PdfReader
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
import requests
from bs4 import BeautifulSoup
# ★追加：スプレッドシートを操作するためのライブラリ
import gspread
from google.oauth2.service_account import Credentials
import datetime
# ★追加：Excelファイルを操作するためのライブラリ
import openpyxl

# ==========================================
# ⚙️ システム設定・マスタ管理（★追加）
# ==========================================
# 今後エージェントが増えたらここに追加するだけでOKです
AGENT_SHEETS = {
    "中倉": "1mPf7VGMYEIN6hYiUWEsFEmDfLNGnx9c4fQM26dhhrM0",
    "福島": "1clnbuoPvHC3yJ9NtWpVGZihi_o5PfQO5JWC_I8h3UCU",
    "木村": "1aJzGK9LMVIjToOTD6Pe4fiGVxV1FXUDaOxY4FqhcIUc",
    "仲本": "1s1whowg_T8IloYB6XrWbK0zEKzurOU1MhwDQFz-TBZI",
    # "山田": "山田用のスプレッドシートID",
}
AGENT_LIST = list(AGENT_SHEETS.keys())

# ==========================================
# 🎨 デザイン定義
# ==========================================
import base64

# 動画を読み込んでBase64形式に変換する関数
def get_base64_video(video_path):
    try:
        with open(video_path, "rb") as video_file:
            return base64.b64encode(video_file.read()).decode()
    except Exception as e:
        return ""

st.set_page_config(page_title="AIエージェントシステム PRO", page_icon="🤖", layout="wide")

# 常にメインのサイバー動画を読み込む
video_base64 = get_base64_video("ScreenRecording_03-04-2026 13-38-53_1.mov")

# 動画が見つかった場合、全画面の背景としてループ再生させるHTMLを挿入
if video_base64:
    video_html = f"""
    <video autoplay loop muted playsinline style="position: fixed; right: 0; bottom: 0; min-width: 100%; min-height: 100%; z-index: -1; object-fit: cover; opacity: 0.8;">
        <source src="data:video/mp4;base64,{video_base64}" type="video/mp4">
    </video>
    """
    st.markdown(video_html, unsafe_allow_html=True)

# ダーク・ネオサイバー仕様のCSS
st.markdown("""
<style>
/* 1. 背景の設定 */
.stApp { background-color: transparent !important; background-image: none !important; }
header { background-color: transparent !important; }

/* 2. メイン画面の文字色 */
.main .block-container h1, .main .block-container h2, .main .block-container h3, 
.main .block-container h4, .main .block-container p, .main .block-container span, 
.main .block-container div {
    color: #FFFFFF !important; text-shadow: 0px 2px 4px rgba(0,0,0,0.8);
}

/* 3. サイドバー */
[data-testid="stSidebar"] {
    background-color: rgba(5, 15, 30, 0.85) !important; backdrop-filter: blur(10px);
}
[data-testid="stSidebar"] p, [data-testid="stSidebar"] span, [data-testid="stSidebar"] h1, 
[data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3, [data-testid="stSidebar"] label {
    color: #00E5FF !important; text-shadow: none !important;
}

/* 4. コンテナ */
.block-container { position: relative; z-index: 1; }

/* 5. ダークグラスモーフィズムパネル */
.cyber-panel, [data-testid="stVerticalBlockBorderWrapper"] {
    background: rgba(10, 20, 40, 0.6) !important; backdrop-filter: blur(8px);
    -webkit-backdrop-filter: blur(8px); border: 1px solid rgba(0, 229, 255, 0.3) !important; 
    box-shadow: 0 8px 32px rgba(0, 0, 0, 0.5) !important; border-radius: 12px !important;
    padding: 25px; margin-top: 20px; position: relative; overflow: hidden;
}

/* 6. HUD装飾 */
.cyber-panel::before, [data-testid="stVerticalBlockBorderWrapper"]::before {
    content: ""; position: absolute; top: 0; left: 0; width: 25px; height: 25px;
    border-top: 3px solid #00E5FF; border-left: 3px solid #00E5FF; border-top-left-radius: 12px;
}
.cyber-panel::after, [data-testid="stVerticalBlockBorderWrapper"]::after {
    content: ""; position: absolute; bottom: 0; right: 0; width: 25px; height: 25px;
    border-bottom: 3px solid #00E5FF; border-right: 3px solid #00E5FF; border-bottom-right-radius: 12px;
}

/* 7. フィードバックボックス */
.fb-box {
    background: rgba(0, 229, 255, 0.1) !important; border-left: 4px solid #00E5FF !important; 
    padding: 15px; margin-top: 10px; border-radius: 6px;
}

/* 8. 入力フォーム */
.stTextInput input, .stTextArea textarea, .stNumberInput input {
    background-color: rgba(0, 0, 0, 0.5) !important; color: #FFFFFF !important;
    border: 1px solid rgba(0, 229, 255, 0.4) !important; border-radius: 6px !important;
}
::placeholder { color: #8899A6 !important; opacity: 1 !important; }

/* 9. メトリック */
[data-testid="stMetricValue"] {
    color: #00E5FF !important; font-weight: bold; font-size: 2.5rem !important; text-shadow: 0 0 10px rgba(0, 229, 255, 0.5);
}

/* 10. ボタン */
.stButton>button { border-radius: 8px !important; font-weight: bold !important; transition: all 0.3s ease; }
.stButton>button[kind="primary"] {
    background: linear-gradient(135deg, #00E5FF 0%, #0077FF 100%) !important; color: white !important;
    border: none !important; box-shadow: 0 4px 10px rgba(0, 229, 255, 0.4) !important;
}
.stButton>button[kind="primary"]:hover {
    box-shadow: 0 6px 15px rgba(0, 229, 255, 0.7) !important; transform: translateY(-2px);
}

/* ★追加 11. ラジオボタンを「サイバーな横タブ」に進化させる魔法のCSS */
div[role="radiogroup"] {
    display: flex; flex-direction: row; justify-content: space-between;
    background: rgba(10, 20, 40, 0.6) !important; padding: 8px; border-radius: 12px;
    border: 1px solid rgba(0, 229, 255, 0.3); box-shadow: 0 4px 15px rgba(0,0,0,0.5);
    margin-bottom: 20px; flex-wrap: wrap; gap: 10px;
}
div[role="radiogroup"] > label {
    flex: 1; text-align: center; background: rgba(255, 255, 255, 0.05) !important;
    border: 1px solid transparent !important; border-radius: 8px !important;
    padding: 12px 10px !important; cursor: pointer; transition: all 0.3s ease;
    display: flex; justify-content: center; align-items: center; min-width: 150px;
}
div[role="radiogroup"] > label:hover {
    background: rgba(0, 229, 255, 0.15) !important;
    border: 1px solid rgba(0, 229, 255, 0.5) !important; transform: translateY(-2px);
}
/* ラジオの丸いアイコンを消してタブ化する */
div[role="radiogroup"] > label > div:first-child { display: none !important; }
div[role="radiogroup"] > label[aria-checked="true"], div[role="radiogroup"] > label[data-checked="true"] {
    background: linear-gradient(135deg, rgba(0, 229, 255, 0.6) 0%, rgba(0, 119, 255, 0.6) 100%) !important;
    border: 1px solid #00E5FF !important; box-shadow: 0 0 20px rgba(0, 229, 255, 0.5);
}
/* ★テキストを強制的に白く表示させる */
div[role="radiogroup"] p, div[role="radiogroup"] span, div[role="radiogroup"] div {
    color: #FFFFFF !important; font-weight: bold !important; margin: 0 !important;
    font-size: 0.95rem !important; text-shadow: 0 1px 3px rgba(0,0,0,0.8);
}
</style>
""", unsafe_allow_html=True)
# ==========================================
# 💾 セッション記憶
# ==========================================
if "history_log" not in st.session_state:
    st.session_state.history_log = [] 
if "carte_log" not in st.session_state:
    st.session_state.carte_log = [] 
if "phase2_generated" not in st.session_state:
    st.session_state.phase2_generated = False 
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []
if "p0_generated" not in st.session_state:
    st.session_state.p0_generated = False
if "p0_interview_date" not in st.session_state:
    st.session_state.p0_interview_date = ""

# ★追加：エラー回避のための初期化
if "p0_change_count" not in st.session_state:
    st.session_state.p0_change_count = ""
if "p0_short_term" not in st.session_state:
    st.session_state.p0_short_term = ""
if "p0_company" not in st.session_state:
    st.session_state.p0_company = ""

# --- セキュリティ ---
LOGIN_PASSWORD = "HR9237"
if "password_correct" not in st.session_state: st.session_state.password_correct = False
if not st.session_state.password_correct:
    st.title("🛡️ システムログイン")
    pwd = st.text_input("アクセスコード", type="password")
    if st.button("ログイン"):
        if pwd == LOGIN_PASSWORD:
            st.session_state.password_correct = True
            st.rerun()
        else: st.error("コードが違います")
    st.stop()

# --- 関数群 ---
def read_files(files):
    content = ""
    for f in files:
        if f.name.endswith('.txt'): content += f.getvalue().decode("utf-8") + "\n"
        elif f.name.endswith('.pdf'):
            try:
                pdf = PdfReader(f)
                for page in pdf.pages: content += (page.extract_text() or "") + "\n"
            except: content += f"[Error: {f.name}]\n"
    return content

def get_url_text(url):
    try:
        res = requests.get(url, timeout=10)
        res.raise_for_status()
        soup = BeautifulSoup(res.content, 'html.parser')
        for script in soup(["script", "style"]):
            script.extract()
        text = soup.get_text(separator='\n')
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return text[:3000] 
    except Exception as e:
        return f"[URL読み取りエラー: {e}]"

def get_section(name, text):
    pattern = f"【{name}】(.*?)(?=【|$)"
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else f"{name}の情報が生成されませんでした。プロンプトを再確認してください。"

def create_docx(history_text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'ＭＳ 明朝'
    style.font._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
    
    doc.add_heading('職務経歴書（自己PR含む）', 0)
    for line in history_text.split('\n'):
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_carte_docx(carte_dict):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'ＭＳ 明朝'
    style.font._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
    
    doc.add_heading('初回面談カルテ', 0)
    for key, value in carte_dict.items():
        doc.add_heading(f'■ {key}', level=2)
        doc.add_paragraph(value)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# ★追加：Excelテンプレートの文字置き換え関数
# ==========================================
def fill_excel_template(template_file, replacement_dict):
    # Excelファイルを読み込む
    wb = openpyxl.load_workbook(template_file)
    
    # 全てのシートを順番にチェック
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                # セルに文字が入っている場合のみ処理
                if cell.value and isinstance(cell.value, str):
                    for key, val in replacement_dict.items():
                        # セルの中に {{志望動機}} などのキーがあれば置き換える
                        if key in cell.value:
                            # 万が一 None などの値が来た時のために str() で文字列化しておく
                            cell.value = cell.value.replace(key, str(val))
                            
    # メモリ上に保存して返す
    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()

# ==========================================
# 📊 スプレッドシート転記・詳細入力メイン関数
# ==========================================
def export_to_spreadsheet(agent_name, seeker_name, interview_date, additional_data=None):
    try:
        credentials_dict = dict(st.secrets["gcp_service_account"])
        scopes = ['https://www.googleapis.com/auth/spreadsheets']
        creds = Credentials.from_service_account_info(credentials_dict, scopes=scopes)
        gc = gspread.authorize(creds)
        
        # ★変更：エージェントごとのシートID振り分け（辞書を使用）
        if agent_name in AGENT_SHEETS:
            sheet_id = AGENT_SHEETS[agent_name]
        else:
            return False, f"登録されていないエージェント名です: {agent_name}"

        sh = gc.open_by_key(sheet_id)
        
        # 1. 原本シートをコピーして個別シート作成
        try:
            original_ws = sh.worksheet("原本")
            new_sheet_name = f"{seeker_name}"
            
            # 同名シートがある場合の重複回避
            existing_sheets = [ws.title for ws in sh.worksheets()]
            if new_sheet_name in existing_sheets:
                new_sheet_name = f"{seeker_name}_{datetime.datetime.now().strftime('%m%d%H%M')}"
            
            new_ws = original_ws.duplicate(insert_sheet_index=1, new_sheet_name=new_sheet_name)
            new_ws_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/edit#gid={new_ws.id}"
            
        except Exception as e:
            return False, f"原本コピー失敗: {e}"

        # 2. 【重要】個別シートの指定セルに情報を入力
        try:
            # A1:B2 結合セルへの名前入力（左上のA1に書き込めばOK）
            new_ws.update_acell('A1', f"{seeker_name} ")
            
            if additional_data:
                # B4: 応募企業名
                new_ws.update_acell('B4', additional_data.get("company_name", ""))
                # D2: 年齢（数字以外を徹底排除）
                raw_age = additional_data.get("age", "")
                # 正規表現で「最初に見つかった連続する数字」だけを抜き出す
                age_match = re.search(r'\d+', raw_age)
                if age_match:
                    age_digits = age_match.group()
                else:
                    age_digits = "" # 数字がない場合は空
                
                new_ws.update_acell('D2', age_digits)
                # E2: 転職回数
                new_ws.update_acell('E2', additional_data.get("change_count", ""))
                # F2: 短期離職数
                new_ws.update_acell('F2', additional_data.get("short_term_leave", ""))
                
                # G2: マネジメント経験 (チェックボックス)
                # 「あり」という文字が含まれていればTrue(チェック)を入れる
                m_exp = additional_data.get("management", "")
                is_m_checked = True if "あり" in m_exp or "経験あり" in m_exp else False
                new_ws.update_acell('G2', is_m_checked)

        except Exception as e:
            st.warning(f"個別シートへの詳細書き込みに一部失敗しました: {e}")

        # 3. 求職者管理表（インデックス）への追記
        try:
            list_ws = sh.worksheet("求職者管理表")
            next_row = len(list_ws.col_values(5)) + 1
            
            final_date = interview_date if interview_date not in ["不明", "記載なし", "なし", ""] else datetime.datetime.now().strftime("%Y/%m/%d")
            hyperlink_formula = f'=HYPERLINK("{new_ws_url}", "{seeker_name}")'
            
            list_ws.update_cell(next_row, 5, hyperlink_formula) # E列: 名前(リンク)
            list_ws.update_cell(next_row, 6, final_date)         # F列: 面談日
            
        except Exception as e:
            return False, f"管理表への追記失敗: {e}"

        return True, f"「{new_sheet_name}」を作成し、データを入力しました！"
        
    except Exception as e:
        return False, f"エラー: {e}"
client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])

# ==========================================
# 画面レイアウト・サイドバー
# ==========================================
st.sidebar.title("AI AGENT MENU")

# ★追加：メイン画面の最上部にタブ風のフェーズ選択を配置（順番も変更！）
app_mode = st.radio(
    "フェーズ選択",
    [
        "1. ランク判定", 
        "2. カルテ作成", 
        "3. 書類作成", 
        "4. 書類審査"
    ],
    horizontal=True,
    label_visibility="collapsed"
)
st.markdown("<br>", unsafe_allow_html=True)

st.sidebar.markdown("---")
agent_name = st.sidebar.text_input("アドバイザー名", placeholder="山田 太郎")
    
# ▼ ここから下を、すべて「一番左の壁」にピッタリ揃えました！ ▼
st.sidebar.divider()
st.sidebar.subheader("📋 面談カルテ履歴 (最新20件)")
if not st.session_state.carte_log:
    st.sidebar.caption("カルテの履歴はありません")
else:
    for i, log in enumerate(st.session_state.carte_log):
        with st.sidebar.expander(f"👤 {log['time']} ({log['name']}様)"):
            if st.button("🔄 復元", key=f"c_res_{i}"):
                # ★追加：履歴復元時に面談日も呼び出す
                st.session_state.p0_interview_date = log["data"].get("面談日", "不明")
                st.session_state.p0_agent = log["data"].get("エージェント名", "")
                st.session_state.p0_seeker = log["data"].get("求職者名", "")
                st.session_state.p0_recog = log["data"].get("エージェント面談の認識", "")
                st.session_state.p0_exp = log["data"].get("エージェントの利用経験", "")
                st.session_state.p0_age = log["data"].get("生年月日・年齢", "")
                st.session_state.p0_cert = log["data"].get("保有資格", "")
                st.session_state.p0_status = log["data"].get("現在の勤務状況", "")
                st.session_state.p0_history = log["data"].get("職務経歴", "")
                st.session_state.p0_reason1 = log["data"].get("転職を考えたきっかけ", "")
                st.session_state.p0_reason2 = log["data"].get("今回の転職で叶えたいこと", "")
                st.session_state.p0_reason3 = log["data"].get("今後のビジョン", "")
                st.session_state.p0_str = log["data"].get("自分の強み", "")
                st.session_state.p0_str_ep = log["data"].get("強みエピソード", "")
                st.session_state.p0_weak = log["data"].get("弱み", "")
                st.session_state.p0_weak_ep = log["data"].get("弱みエピソード", "")
                st.session_state.p0_c_job = log["data"].get("希望職種・業務", "")
                st.session_state.p0_c_loc = log["data"].get("希望勤務地", "")
                st.session_state.p0_c_cur_sal = log["data"].get("現在年収・給与", "")
                st.session_state.p0_c_req_sal = log["data"].get("希望年収・給与", "")
                st.session_state.p0_c_time = log["data"].get("勤務時間・休日", "")
                st.session_state.p0_c_vibes = log["data"].get("社風・雰囲気", "")
                st.session_state.p0_c_date = log["data"].get("入社希望日", "")
                st.session_state.p0_o_ans = log["data"].get("確認事項や不安ごと", "")
                st.session_state.p0_o_ndate = log["data"].get("次回面談日", "")
                st.session_state.p0_o_ntime = log["data"].get("次回面談時間", "")
                st.session_state.p0_generated = True
                st.rerun()
                
            dl_doc_c = create_carte_docx(log["data"])
            st.download_button(
                label="📥 WordをDL",
                data=dl_doc_c,
                file_name=f"履歴_面談カルテ_{log['name']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"hist_dl_c_{i}"
            )

st.sidebar.divider()
st.sidebar.subheader("📄 書類生成履歴 (最新20件)")
if not st.session_state.history_log:
    st.sidebar.caption("書類履歴はありません")
else:
    for i, log in enumerate(st.session_state.history_log):
        with st.sidebar.expander(f"📁 {log['time']} ({log['job']})"):
            if st.button("🔄 この画面を復元する", key=f"restore_btn_{i}"):
                st.session_state.phase2_score = log["score"]
                st.session_state.phase2_advice = log["advice"]
                st.session_state.phase2_combined = log["combined"]
                st.session_state.phase2_motive = log.get("motive", "")
                st.session_state.chat_messages = log["chat"]
                st.session_state.phase2_generated = True
                st.rerun()
                
            dl_doc = create_docx(log["combined"])
            st.download_button(
                label="📥 WordをDL",
                data=dl_doc,
                file_name=f"履歴_職務経歴書_{i}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"hist_dl_{i}"
            )

# ==========================================
# Phase 0: 初回面談 (カルテ作成)
# ==========================================
if app_mode == "2. カルテ作成":
    st.title("初回面談時：カルテ自動生成")
    st.markdown("文字起こしファイル(PDF/TXT)を添付するか、テキストを直接貼り付けてください。AIが自動で項目を整理します。")

    components.html("""
    <div style="font-family: sans-serif; margin-bottom: 10px;">
        <p style="color: #00E5FF; font-size: 14px; font-weight: bold; margin-bottom: 5px;">🎤 音声入力（補助ツール）</p>
        <button id="start-btn" style="background: transparent; color: #00E5FF; border: 1px solid #00E5FF; border-radius: 5px; padding: 5px 10px; cursor: pointer;">🔴 録音開始</button>
        <button id="stop-btn" style="background: transparent; color: #ff4b4b; border: 1px solid #ff4b4b; border-radius: 5px; padding: 5px 10px; cursor: pointer;" disabled>⏹ 停止</button>
        <p style="color: #FFFFFF; font-size: 12px; margin-top: 5px;">※録音した場合は下のテキストエリアに自動で入力されます</p>
    </div>
    <script>
        const startBtn = document.getElementById('start-btn'); const stopBtn = document.getElementById('stop-btn');
        let recognition;
        if ('webkitSpeechRecognition' in window) {
            recognition = new webkitSpeechRecognition(); recognition.lang = 'ja-JP'; recognition.continuous = true;
            recognition.onresult = function(event) {
                let finalTranscript = '';
                for (let i = event.resultIndex; i < event.results.length; ++i) {
                    if (event.results[i].isFinal) finalTranscript += event.results[i][0].transcript;
                }
            };
            startBtn.onclick = () => { recognition.start(); startBtn.disabled = true; stopBtn.disabled = false; };
            stopBtn.onclick = () => { recognition.stop(); startBtn.disabled = false; stopBtn.disabled = true; };
        }
    </script>
    """, height=90)

    u_files_memo = st.file_uploader("📂 文字起こしファイルなど (PDF/TXT)", accept_multiple_files=True, key="p0_up")
    raw_memo = st.text_area("📝 面談メモ / 文字起こしテキスト (手入力・コピペ用)", height=200, placeholder="ここにテキストをペースト、または手書きメモを入力してください...")

    if st.button("🪄 AIで項目を自動抽出", type="primary"):
        file_text = read_files(u_files_memo) if u_files_memo else ""
        combined_memo = file_text + "\n" + raw_memo

        if not combined_memo.strip():
            st.warning("文字起こしファイルを添付するか、メモを入力してください。")
        else:
            with st.spinner("AIが面談内容を詳細に分析中..."):
                # ★変更：プロンプトを大幅強化（エージェント名選択式、職歴8項目指定）
                prompt = f"""
                あなたは優秀なキャリアアドバイザーのアシスタントです。
                以下の「面談の文字起こし・メモ」から、求職者の情報を抽出して整理してください。
                情報が語られていない項目は「不明」または「記載なし」と記載してください。

                【面談データ】
                {combined_memo}

                【抽出フォーマット（絶対厳守）】
                以下の【】で囲まれたセクション名を必ず使用し、各項目を個別に抽出してください。

                【面談日】
                (YYYY/MM/DD形式)
                【エージェント名】
                (必ず以下のリストから完全一致で選択してください。該当なしは「その他」：{AGENT_LIST})
                【求職者名】
                【エージェント面談の認識】
                【エージェントの利用経験】
                【生年月日・年齢】
                【保有資格】
                【現在の勤務状況】

                # --- 今回追加した重要項目 ---
                【転職回数】
                (在職中も含めた合計社数-1)
                【短期離職数】
                (1年以内の離職回数)
                【応募企業名】
                (具体的な社名があれば記載、なければ「（未入力）」)
                # --------------------------

                【職務経歴】
                (経験社数分、以下の8項目を必ず箇条書きで詳細に抽出すること)
                ■会社名：
                ・雇用形態：(正社員、アルバイト、業務委託、契約社員など)
                ・部署／役職：(あれば記載、なければなし)
                ・職種：
                ・主な業務内容：
                ・入社理由：
                ・実績や成果：
                ・退職理由：

                【転職を考えたきっかけ】
                【今回の転職で叶えたいこと】
                【入社後どうなっていたいか】
                【自分の強み】
                【強みエピソード】
                【弱み】
                【自分の弱みエピソード】
                【希望職種・業務】
                【希望勤務地】
                【現在年収・給与】
                【希望年収・給与】
                【勤務時間・休日】
                【社風・雰囲気】
                【入社希望日】
                【求職者からの確認事項や不安ごと】
                【次回面談日】
                【次回面談時間】
                """
                try:
                    resp = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    res = resp.text

                    st.session_state.p0_interview_date = get_section("面談日", res)
                    st.session_state.p0_agent = get_section("エージェント名", res)
                    st.session_state.p0_seeker = get_section("求職者名", res)
                    
                    # ★追加項目の受け取り
                    st.session_state.p0_change_count = get_section("転職回数", res)
                    st.session_state.p0_short_term = get_section("短期離職数", res)
                    st.session_state.p0_company = get_section("応募企業名", res)

                    # (以下、既存の get_section 処理を継続...)
                    st.session_state.p0_recog = get_section("エージェント面談の認識", res)
                    st.session_state.p0_exp = get_section("エージェントの利用経験", res)
                    st.session_state.p0_age = get_section("生年月日・年齢", res)
                    st.session_state.p0_cert = get_section("保有資格", res)
                    st.session_state.p0_status = get_section("現在の勤務状況", res)
                    st.session_state.p0_history = get_section("職務経歴", res)
                    st.session_state.p0_reason1 = get_section("転職を考えたきっかけ", res)
                    st.session_state.p0_reason2 = get_section("今回の転職で叶えたいこと", res)
                    st.session_state.p0_reason3 = get_section("入社後どうなっていたいか", res)
                    st.session_state.p0_str = get_section("自分の強み", res)
                    st.session_state.p0_str_ep = get_section("強みエピソード", res)
                    st.session_state.p0_weak = get_section("弱み", res)
                    st.session_state.p0_weak_ep = get_section("自分の弱みエピソード", res)
                    st.session_state.p0_c_job = get_section("希望職種・業務", res)
                    st.session_state.p0_c_loc = get_section("希望勤務地", res)
                    st.session_state.p0_c_cur_sal = get_section("現在年収・給与", res)
                    st.session_state.p0_c_req_sal = get_section("希望年収・給与", res)
                    st.session_state.p0_c_time = get_section("勤務時間・休日", res)
                    st.session_state.p0_c_vibes = get_section("社風・雰囲気", res)
                    st.session_state.p0_c_date = get_section("入社希望日", res)
                    st.session_state.p0_o_ans = get_section("求職者からの確認事項や不安ごと", res)
                    st.session_state.p0_o_ndate = get_section("次回面談日", res)
                    st.session_state.p0_o_ntime = get_section("次回面談時間", res)
                    
                    st.session_state.p0_generated = True

                    carte_dict = {
                        "面談日": st.session_state.p0_interview_date, # ★追加
                        "エージェント名": st.session_state.p0_agent, "求職者名": st.session_state.p0_seeker,
                        "エージェント面談の認識": st.session_state.p0_recog, "エージェントの利用経験": st.session_state.p0_exp,
                        "生年月日・年齢": st.session_state.p0_age, "保有資格": st.session_state.p0_cert, "現在の勤務状況": st.session_state.p0_status,
                        "職務経歴": st.session_state.p0_history,
                        "転職を考えたきっかけ": st.session_state.p0_reason1, "今回の転職で叶えたいこと": st.session_state.p0_reason2, "今後のビジョン": st.session_state.p0_reason3,
                        "自分の強み": st.session_state.p0_str, "強みエピソード": st.session_state.p0_str_ep, "弱み": st.session_state.p0_weak, "弱みエピソード": st.session_state.p0_weak_ep,
                        "希望職種・業務": st.session_state.p0_c_job, "希望勤務地": st.session_state.p0_c_loc, "現在年収・給与": st.session_state.p0_c_cur_sal, "希望年収・給与": st.session_state.p0_c_req_sal,
                        "勤務時間・休日": st.session_state.p0_c_time, "社風・雰囲気": st.session_state.p0_c_vibes, "入社希望日": st.session_state.p0_c_date,
                        "確認事項や不安ごと": st.session_state.p0_o_ans, "次回面談日": st.session_state.p0_o_ndate, "次回面談時間": st.session_state.p0_o_ntime
                    }
                    st.session_state.carte_log.insert(0, {
                        "time": time.strftime('%Y/%m/%d %H:%M'),
                        "name": st.session_state.p0_seeker if st.session_state.p0_seeker else "未入力",
                        "data": carte_dict
                    })
                    if len(st.session_state.carte_log) > 20: st.session_state.carte_log.pop()

                except Exception as e:
                    st.error(f"解析エラー: {e}")

    # 自動抽出されたデータの表示と編集
    if st.session_state.get("p0_generated"):
        st.markdown(f'<div class="cyber-panel"><div class="scan-line"></div><h3>📋 抽出されたカルテ情報</h3><p style="color:white; font-size:14px;">※手作業で修正・追記が可能です</p></div>', unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        st.subheader("📄 職務経歴書に直結する情報")
        with st.container(border=True):
            st.markdown('<div class="emerald-box"></div>', unsafe_allow_html=True)
            e_seeker = st.text_input("求職者名", value=st.session_state.p0_seeker)
            
            st.markdown("#### 🏢 職務経歴")
            e_history = st.text_area("職務経歴 (複数社対応)", value=st.session_state.p0_history, height=250)
            
            st.markdown("#### 🚀 転職理由・キャリア観")
            c4, c5, c6 = st.columns(3)
            with c4: e_reason1 = st.text_area("転職を考えたきっかけ", value=st.session_state.p0_reason1, height=120)
            with c5: e_reason2 = st.text_area("転職で叶えたいこと", value=st.session_state.p0_reason2, height=120)
            with c6: e_reason3 = st.text_area("今後のビジョン", value=st.session_state.p0_reason3, height=120)

            st.markdown("#### 💪 強み・弱み")
            c7, c8 = st.columns(2)
            with c7:
                e_str = st.text_input("自分の強み", value=st.session_state.p0_str)
                e_str_ep = st.text_area("強みエピソード", value=st.session_state.p0_str_ep, height=100)
            with c8:
                e_weak = st.text_input("弱み", value=st.session_state.p0_weak)
                e_weak_ep = st.text_area("弱みエピソード", value=st.session_state.p0_weak_ep, height=100)

        st.markdown("<br>", unsafe_allow_html=True)
        st.subheader("🏢 エージェント管理・条件情報")
        with st.container(border=True):
            st.markdown('<div class="emerald-box"></div>', unsafe_allow_html=True)
            
            # ★変更：エージェント名をリストから選べるようにセレクトボックス化
            c_ag1, c_ag2 = st.columns(2)
            with c_ag1:
                agent_idx = AGENT_LIST.index(st.session_state.p0_agent) if st.session_state.p0_agent in AGENT_LIST else 0
                e_agent = st.selectbox("エージェント名", AGENT_LIST + ["その他"], index=agent_idx)
            with c_ag2: e_interview_date = st.text_input("面談日 (不明・空欄時は今日の日付で転記)", value=st.session_state.p0_interview_date)
            
            st.markdown("#### 👤 基本情報")
            c1, c2, c3 = st.columns(3)
            with c1:
                e_status = st.text_input("現在の勤務状況", value=st.session_state.p0_status)
                e_cert = st.text_input("保有資格", value=st.session_state.p0_cert)
            with c2:
                e_recog = st.text_input("面談の認識(有/無)", value=st.session_state.p0_recog)
                e_exp = st.text_input("利用経験(有/無)", value=st.session_state.p0_exp)
            with c3:
                e_age = st.text_input("生年月日・年齢", value=st.session_state.p0_age)
            
            st.markdown("#### 🎯 就職活動希望条件")
            c9, c10, c11 = st.columns(3)
            with c9:
                e_c_job = st.text_input("希望職種・業務", value=st.session_state.p0_c_job)
                e_company = st.text_input("応募企業名", value=st.session_state.p0_company) # 追加
                e_c_loc = st.text_input("希望勤務地", value=st.session_state.p0_c_loc)
                e_c_date = st.text_input("入社希望日", value=st.session_state.p0_c_date)
            with c10:
                e_c_cur_sal = st.text_input("現在年収・給与", value=st.session_state.p0_c_cur_sal)
                e_c_req_sal = st.text_input("希望年収・給与", value=st.session_state.p0_c_req_sal)
                e_change_count = st.text_input("転職回数", value=st.session_state.p0_change_count) # 追加
                e_short_term = st.text_input("短期離職数", value=st.session_state.p0_short_term)   # 追加
            with c11:
                e_c_time = st.text_input("勤務時間・休日", value=st.session_state.p0_c_time)
                e_c_vibes = st.text_input("社風・雰囲気", value=st.session_state.p0_c_vibes)

            st.markdown("#### 📅 その他確認・次回設定")
            c12, c13 = st.columns([2, 1])
            with c12: e_o_ans = st.text_area("確認事項や不安ごと", value=st.session_state.p0_o_ans, height=100)
            with c13:
                e_o_ndate = st.text_input("次回面談日", value=st.session_state.p0_o_ndate)
                e_o_ntime = st.text_input("次回面談時間", value=st.session_state.p0_o_ntime)

        # 出力ボタン群
        st.divider()
        c_btn_w, c_btn_s, _ = st.columns([1, 1, 2])
        
        with c_btn_w:
            carte_dict_updated = {
                "面談日": e_interview_date, # ★追加
                "エージェント名": e_agent, "求職者名": e_seeker,
                "エージェント面談の認識": e_recog, "エージェントの利用経験": e_exp,
                "生年月日・年齢": e_age, "保有資格": e_cert, "現在の勤務状況": e_status,
                "職務経歴": e_history,
                "転職を考えたきっかけ": e_reason1, "今回の転職で叶えたいこと": e_reason2, "今後のビジョン": e_reason3,
                "自分の強み": e_str, "強みエピソード": e_str_ep, "弱み": e_weak, "弱みエピソード": e_weak_ep,
                "希望職種・業務": e_c_job, "希望勤務地": e_c_loc, "現在年収・給与": e_c_cur_sal, "希望年収・給与": e_c_req_sal,
                "勤務時間・休日": e_c_time, "社風・雰囲気": e_c_vibes, "入社希望日": e_c_date,
                "確認事項や不安ごと": e_o_ans, "次回面談日": e_o_ndate, "次回面談時間": e_o_ntime
            }
            docx_file = create_carte_docx(carte_dict_updated)
            st.download_button(
                label="📥 この面談カルテをWordでDL",
                data=docx_file,
                file_name=f"面談カルテ_{e_seeker}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )

        with c_btn_s:
            if st.button("📊 スプレッドシートに自動転記", type="primary", use_container_width=True):
                with st.spinner("スプレッドシートを作成中..."):
                    # 画面上で修正した最新の値をセットして送信
                    info = {
                        "company_name": e_company,      # 修正済み
                        "age": e_age,
                        "change_count": e_change_count, # 修正済み
                        "short_term_leave": e_short_term, # 修正済み
                        "management": e_history
                    }

                    success, message = export_to_spreadsheet(e_agent, e_seeker, e_interview_date, info)
                    
                    if success:
                        st.success(message)
                    else:
                        st.error(message)

# ==========================================
# Phase 1: 応募時 (ランク判定)
# ==========================================
elif app_mode == "1. ランク判定":
    st.title("応募時簡易分析: ランク判定")
    col1, col2, col3 = st.columns(3)
    with col1: age = st.number_input("年齢", 18, 85, 25) 
    with col2: job_changes = st.number_input("転職回数", 0, 15, 1)
    with col3: short_term = st.number_input("短期離職数", 0, 10, 0)
    
    if st.button("ランクを判定する", type="primary"):
        if age < 20: age_s = -8
        elif 20 <= age <= 21: age_s = 8
        elif 22 <= age <= 25: age_s = 10
        elif 26 <= age <= 29: age_s = 8
        elif 30 <= age <= 35: age_s = 7
        else: age_s = 5

        job_bonus = 0
        if age <= 24 and job_changes == 0: job_bonus = 10
        elif 25 <= age <= 29 and job_changes <= 1: job_bonus = 10
        elif 25 <= age <= 29 and job_changes <= 2: job_bonus = 7
        elif 30 <= age <= 35 and job_changes <= 2: job_bonus = 10
        elif 30 <= age <= 35 and job_changes <= 3: job_bonus = 7
        elif 35 <= age <= 85 and job_changes <= 2: job_bonus = 10
        elif 35 <= age <= 85 and job_changes <= 3: job_bonus = 7
        elif 50 <= age <= 85 and job_changes <= 4: job_bonus = 5
        elif job_changes <= 1: job_bonus = 5

        job_penalty = 0
        if job_changes == 2: job_penalty = -5
        elif job_changes == 3: job_penalty = -10
        elif job_changes >= 5: job_penalty = -20
        
        st_penalty = short_term * 10
        total = age_s + job_bonus + job_penalty - st_penalty + 5 

        if total >= 23: cn, rc = "優秀 (Class-S)", "#00ff00"
        elif total >= 18: cn, rc = "良好 (Class-A)", "#00e5ff"
        elif total >= 13: cn, rc = "標準 (Class-B)", "#ffff00"
        elif total >= 8: cn, rc = "要努力 (Class-C)", "#ff9900"
        else: cn, rc = "測定不能 (Class-Z)", "#ff0000"

        st.markdown(f'<div class="cyber-panel"><h3>判定結果: <span style="color:{rc};">{cn}</span></h3></div>', unsafe_allow_html=True)
        if total >= 15: st.success("NICE❕ **【エージェント指示】** 優先度：高")
        elif 7 <= total < 15: st.info("safe **【エージェント指示】** 優先度：中")
        else: st.error("🚨 **【エージェント指示】** 優先度：低")

# ==========================================
# Phase 2: 初回面談後 (詳細分析/書類作成)
# ==========================================
elif app_mode == "3. 書類作成":
    st.title("初回面談後: 書類作成＆分析")
    
    c_top1, c_top2 = st.columns(2)
    with c_top1: t_ind = st.text_input("志望業種", placeholder="未入力の場合は添付資料から判断します")
    with c_top2: t_job = st.text_input("志望職種", placeholder="未入力の場合は添付資料から判断します")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("🏢 企業・募集情報")
        u_url_corp = st.text_input("🔗 求人票URL (自動読み取り)", placeholder="https://...")
        u_files_corp = st.file_uploader("企業求人票など", accept_multiple_files=True, key="corp_up")
        
    with col2:
        st.subheader("📂 求職者情報")
        
        # ★追加：Phase 0のカルテ情報を読み込むボタン
        if st.button("🔄 Phase 0のカルテ情報を読み込む"):
            st.session_state.p2_sync_achievement = f"【職務経歴】\n{st.session_state.p0_history}\n\n【転職理由】\n{st.session_state.p0_reason1}\n\n【叶えたいこと】\n{st.session_state.p0_reason2}\n\n【強み】\n{st.session_state.p0_str}\n{st.session_state.p0_str_ep}"
            st.success("Phase 0のカルテデータを読み込みました！")

        u_files_seeker = st.file_uploader("履歴書・面談文字起こし", accept_multiple_files=True, key="seeker_up")
        
        # ★変更：同期したデータを初期値としてセット
        achievement = st.text_area("求職者の補足事項・メモ", value=st.session_state.get("p2_sync_achievement", ""), height=100)
        
        components.html("""
        <div style="font-family: sans-serif; margin-top: -10px;">
            <p style="color: #00E5FF; font-size: 14px; font-weight: bold; margin-bottom: 5px;">🎤 音声入力</p>
            <button id="start-btn" style="background: transparent; color: #00E5FF; border: 1px solid #00E5FF; border-radius: 5px; padding: 5px 10px; cursor: pointer;">🔴 録音開始</button>
            <button id="stop-btn" style="background: transparent; color: #ff4b4b; border: 1px solid #ff4b4b; border-radius: 5px; padding: 5px 10px; cursor: pointer;" disabled>⏹ 停止</button>
            <textarea id="result" style="width: 100%; height: 70px; background: rgba(0,0,0,0.3); color: white; border: 1px solid #00E5FF; border-radius: 5px; padding: 5px; margin-top: 5px;"></textarea>
        </div>
        <script>
            const startBtn = document.getElementById('start-btn'); const stopBtn = document.getElementById('stop-btn');
            const resultArea = document.getElementById('result'); let recognition;
            if ('webkitSpeechRecognition' in window) {
                recognition = new webkitSpeechRecognition(); recognition.lang = 'ja-JP'; recognition.continuous = true;
                recognition.onresult = function(event) {
                    let finalTranscript = '';
                    for (let i = event.resultIndex; i < event.results.length; ++i) {
                        if (event.results[i].isFinal) finalTranscript += event.results[i][0].transcript;
                    }
                    if(finalTranscript) resultArea.value += finalTranscript + '\\n';
                };
                startBtn.onclick = () => { recognition.start(); startBtn.disabled = true; stopBtn.disabled = false; };
                stopBtn.onclick = () => { recognition.stop(); startBtn.disabled = false; stopBtn.disabled = true; };
            }
        </script>
        """, height=160)

    if st.button("AI書類生成を開始", type="primary"):
        corp_url_data = get_url_text(u_url_corp) if u_url_corp else ""
        corp_file_data = read_files(u_files_corp) if u_files_corp else ""
        corp_data = corp_file_data + "\n" + corp_url_data
        seeker_data = read_files(u_files_seeker) if u_files_seeker else ""
        
        if not (t_ind or t_job or corp_data.strip()): st.warning("企業情報を入力してください。")
        elif not (achievement or seeker_data.strip()): st.warning("求職者情報を入力してください。")
        else:
            with st.spinner("情報を深く分析中..."):
                prompt = f"""
あなたは人材紹介会社の**プロキャリアライター兼採用目線の職務経歴書編集者**です。
提供された「企業情報」と「求職者情報」を深く分析し、企業が「ぜひ会ってみたい」と思える具体的・誠実・読みやすい書類を作成してください。

【入力：企業情報】
志望業種：{t_ind if t_ind else "未入力（企業資料から判断してください）"}
志望職種：{t_job if t_job else "未入力（企業資料から判断してください）"}
企業側資料：{corp_data if corp_data else "なし"}

【入力：求職者情報】
補足メモ：{achievement if achievement else "なし"}
求職者側資料（履歴書・面談文字起こし等）：{seeker_data if seeker_data else "なし"}

---
【重要ルール】
- 提供された「求職者情報」から、実際の経験・業務内容・成果を具体的に抽出し、必ず書類に反映させてください。架空の経験は絶対に書かないでください。
- 「企業情報」の求める人物像に合わせ、求職者の強みを最適化して記載してください。
- 以下の【】で囲まれた各セクションを、「一切省略せずに」出力してください。

【評価】
(S最高/A良き！/Bいい感じ/C要努力/Z測定不能のみ)

【理由とアドバイス】
(評価の理由と、プロのエージェントとしての鋭い視点や、面接での深掘りポイントのアドバイスを記載)

【職務経歴】
1. 作成日・氏名
2. 職務経歴（各社ごとに以下の構成を維持。必ず求職者資料の事実を元に書くこと）
   ■会社名：〇〇
   雇用形態：〇〇
   事業内容：〇〇
   役職：〇〇
   ▼業務内容
   ・実際の業務内容（タスク・役割）を5〜7行で具体的に記載。
   ・【絶対ルール】文末は「〜ました」「〜です」を禁止し、必ず「〇〇を実施。」「〇〇を担当。」「〇〇に貢献。」と簡潔に言い切ること。
   ▼成果
   ・数値・改善・貢献を具体的に記載。
   ・【絶対ルール】文末は「〜ました」を禁止し、「〇〇を実現し、〇〇％改善。」「〇〇件の契約を継続的に達成。」のように言い切ること。

【自己PR】
- 企業情報に最適化された自己PR。
- 企業の理念・社風・仕事内容に合わせ、経験をどう活かせるか、なぜ惹かれたかを記載。
- 400字で構成。事実を元にし、嘘や推測は含めない。
- 【絶対ルール】長文で読みにくくなるのを防ぐため、内容の区切りごとに必ず「改行（段落分け）」を行ってください。
- 「」や””や・などAI文章だとわかる記号は控える。文体は敬体（です・ます）。
- 一文は60文字以内で簡潔に。丁寧・誠実・安定感のある文体で統一。
- キャリアコンサルタントの視点を取り入れ、求職者の泥臭い努力や具体的なエピソードを魅力的に引き立たせてください。

【志望動機】
- 企業情報と求職者情報を結びつけ、なぜこの企業なのかを具体的に記載。
- ありきたりな言葉ではなく、その企業ならではの強みや特徴に惹かれた理由を深く掘り下げてください。
- 企業にマイナスにならないのを前提に、年齢に合わせた文章・言葉使いにする。
- 約450字で作成。業務や実績は推測や嘘を避ける。
- 【絶対ルール】長文で読みにくくなるのを防ぐため、内容の区切りごとに必ず「改行（段落分け）」を行ってください。
- 「」や””や・などは控える。
"""
                try:
                    resp = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    res = resp.text
                    st.session_state.phase2_score = get_section("評価", res)
                    st.session_state.phase2_advice = get_section("理由とアドバイス", res)
                    hist = get_section('職務経歴', res)
                    pr = get_section('自己PR', res)
                    st.session_state.phase2_combined = f"{hist}\n\n■自己PR\n{pr}"
                    st.session_state.phase2_motive = get_section('志望動機', res)
                    st.session_state.phase2_generated = True
                    st.session_state.chat_messages = [] 
                    
                    st.session_state.history_log.insert(0, {
                        "time": time.strftime('%Y/%m/%d %H:%M'),
                        "job": t_job if t_job else "未指定",
                        "score": st.session_state.phase2_score,
                        "advice": st.session_state.phase2_advice,
                        "combined": st.session_state.phase2_combined,
                        "motive": st.session_state.phase2_motive,
                        "chat": []
                    })
                    if len(st.session_state.history_log) > 20: st.session_state.history_log.pop()
                except Exception as e: st.error(f"解析エラー: {e}")

    if st.session_state.get("phase2_generated"):
        st.markdown(f'<div class="cyber-panel"><div class="scan-line"></div><h3>AI分析評価スコア: {st.session_state.phase2_score}</h3><div class="fb-box">{st.session_state.phase2_advice}</div></div>', unsafe_allow_html=True)
        st.divider()
        st.subheader("📄 職務経歴書（自己PR含む・高品質版）")
        st.code(st.session_state.phase2_combined, language="text")
        
        c_btn1, c_btn2, _ = st.columns([1.5, 1.2, 3])
        with c_btn1:
            docx_file = create_docx(st.session_state.phase2_combined)
            st.download_button(label="📥 WordでDL", data=docx_file, file_name=f"職務経歴書_{time.strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with c_btn2:
            components.html("""<button onclick="window.parent.print()" style="background:transparent; color:#00E5FF; border:1px solid #00E5FF; padding:8px 12px; border-radius:8px; font-size:13px; cursor:pointer; width:auto;">🖨️ PDF保存</button>""", height=50)
        
        st.subheader("📄 志望動機")
        st.code(st.session_state.phase2_motive, language="text")

        # ==========================================
        # ★追加：デザイン済みExcel履歴書への完全自動流し込み機能
        # ==========================================
        st.divider()
        st.subheader("📊 Excel履歴書への完全自動流し込み")
        st.info("※タグ（{{ふりがな}}や{{歴年1}}など）を配置したExcelテンプレートをアップロードしてください。")
        
        u_excel_template = st.file_uploader("履歴書テンプレート (Excel形式: .xlsx)", type=["xlsx"], key="excel_tpl_up")
        
        if u_excel_template:
            if st.button("✨ 履歴書データを解析してExcelに流し込む", type="primary"):
                with st.spinner("求職者のデータを解析し、Excelにマッピング中..."):
                    try:
                        import json
                        seeker_raw_data = read_files(u_files_seeker) if u_files_seeker else ""
                        seeker_raw_data += "\n" + achievement
                        
                        json_prompt = f"""
                        あなたは履歴書データ抽出のプロです。以下の求職者データから、履歴書作成に必要な項目を抽出し、必ず以下のJSONフォーマット（プレーンテキスト）のみを出力してください。マークダウン（```jsonなど）は不要です。

                        【抽出ルール】
                        - データが存在しない項目は空文字("")にしてください。
                        - 年はすべて「西暦」で統一してください。
                        - 学歴と職歴は同じリスト(history)にまとめ、学歴の先頭には {{"year":"", "month":"", "content":"学歴"}} を、職歴の先頭には {{"year":"", "month":"", "content":"職歴"}} となる見出し行を必ず挿入してください。
                        - 資格も同様にリスト(licenses)にしてください。

                        【求職者データ】
                        {seeker_raw_data}

                        【出力フォーマット】
                        {{
                          "furigana": "ふりがな",
                          "name": "氏名",
                          "gender": "男/女",
                          "birth_age": "199X年X月X日生 (満XX歳)",
                          "zip_code": "〒XXX-XXXX",
                          "address_furigana": "じゅうしょふりがな",
                          "address": "住所",
                          "phone": "電話番号",
                          "email": "メールアドレス",
                          "history": [
                             {{"year": "", "month": "", "content": "学歴"}},
                             {{"year": "2015", "month": "4", "content": "〇〇高校 入学"}},
                             {{"year": "2018", "month": "3", "content": "〇〇高校 卒業"}},
                             {{"year": "", "month": "", "content": "職歴"}},
                             {{"year": "2018", "month": "4", "content": "株式会社〇〇 入社"}}
                          ],
                          "licenses": [
                             {{"year": "2020", "month": "10", "content": "TOEIC公開テスト 800点取得"}}
                          ]
                        }}
                        """
                        # AIにJSONを出力させる
                        resp_json = client.models.generate_content(model='gemini-2.5-flash', contents=json_prompt)
                        # JSON部分だけをクリーンに取り出す
                        json_text = resp_json.text.replace('```json', '').replace('```', '').strip()
                        data = json.loads(json_text)
                        
                        # 置き換え用辞書の作成
                        replacements = {
                            "{{ふりがな}}": data.get("furigana", ""),
                            "{{氏名}}": data.get("name", ""),
                            "{{性別}}": data.get("gender", ""),
                            "{{生年月日_年齢}}": data.get("birth_age", ""),
                            "{{郵便番号}}": data.get("zip_code", ""),
                            "{{住所ふりがな}}": data.get("address_furigana", ""),
                            "{{住所}}": data.get("address", ""),
                            "{{電話番号}}": data.get("phone", ""),
                            "{{Email}}": data.get("email", ""),
                            "{{志望動機}}": st.session_state.phase2_motive,
                        }

                        # 学歴・職歴の流し込み（最大30行まで対応。余ったタグは空白化）
                        history = data.get("history", [])
                        for i in range(1, 31):
                            if i <= len(history):
                                replacements[f"{{{{歴年{i}}}}}"] = history[i-1].get("year", "")
                                replacements[f"{{{{歴月{i}}}}}"] = history[i-1].get("month", "")
                                replacements[f"{{{{歴内容{i}}}}}"] = history[i-1].get("content", "")
                            else:
                                replacements[f"{{{{歴年{i}}}}}"] = ""
                                replacements[f"{{{{歴月{i}}}}}"] = ""
                                replacements[f"{{{{歴内容{i}}}}}"] = ""

                        # 資格の流し込み（最大10行まで対応。余ったタグは空白化）
                        licenses = data.get("licenses", [])
                        for i in range(1, 11):
                            if i <= len(licenses):
                                replacements[f"{{{{資格年{i}}}}}"] = licenses[i-1].get("year", "")
                                replacements[f"{{{{資格月{i}}}}}"] = licenses[i-1].get("month", "")
                                replacements[f"{{{{資格内容{i}}}}}"] = licenses[i-1].get("content", "")
                            else:
                                replacements[f"{{{{資格年{i}}}}}"] = ""
                                replacements[f"{{{{資格月{i}}}}}"] = ""
                                replacements[f"{{{{資格内容{i}}}}}"] = ""

                        # 先ほど作った関数でExcelを書き換え
                        new_excel_data = fill_excel_template(u_excel_template, replacements)
                        
                        st.success("✨ データの抽出とExcelへの流し込みが完了しました！")
                        safe_seeker_name = st.session_state.p0_seeker if st.session_state.get("p0_seeker") else "求職者"
                        st.download_button(
                            label="📥 完成したExcel履歴書をダウンロード",
                            data=new_excel_data,
                            file_name=f"完成版_履歴書_{safe_seeker_name}_{time.strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            type="primary"
                        )
                        
                    except json.JSONDecodeError:
                        st.error("AIからのデータ抽出に失敗しました。もう一度ボタンを押してお試しください。")
                    except Exception as e:
                        st.error(f"エラーが発生しました: {e}")

        # --- AIチャット機能 ---
        st.divider()
        st.subheader("💬 AIアシスタントと内容を調整する")
        
        edit_target = st.radio("🎯 修正する項目を選択", ["全体", "職務経歴", "自己PR", "志望動機"], horizontal=True)

        for msg in st.session_state.chat_messages:
            with st.chat_message(msg["role"]): st.markdown(msg["content"])
                
        if chat_input := st.chat_input(f"【{edit_target}】への修正依頼を入力してください"):
            st.session_state.chat_messages.append({"role": "user", "content": f"[{edit_target}] {chat_input}"})
            with st.chat_message("user"): st.markdown(f"**[{edit_target}]** {chat_input}")
                
            with st.chat_message("assistant"):
                chat_prompt = f"""
あなたはプロのキャリアコンサルタントです。ユーザーの【修正指示】に基づき、書類をより魅力的で具体的に改善してください。

【対象セクション】: {edit_target}

【厳守ルール】
1. 指定された「{edit_target}」の部分のみを修正してください（「全体」の場合は全体のバランスを見て修正）。
2. 修正対象外のセクションは一切変更せず、そのまま出力してください。
3. 元の改行、見出し(■,▼,・)、箇条書き、体言止めのフォーマットを絶対に崩さない。
4. 自己PRや志望動機を修正する場合は、必ず適度な「改行（段落分け）」を入れて読みやすくすること。
5. プロの視点から、より具体的で説得力のある表現にブラッシュアップすること。

【現在の書類データ】
{st.session_state.phase2_combined}
志望動機：{st.session_state.phase2_motive}

【ユーザーからの修正指示】
{chat_input}
"""
                try:
                    chat_resp = client.models.generate_content(model='gemini-2.5-flash', contents=chat_prompt)
                    st.markdown(chat_resp.text)
                    st.session_state.chat_messages.append({"role": "assistant", "content": chat_resp.text})
                    if st.session_state.history_log:
                        st.session_state.history_log[0]["chat"] = st.session_state.chat_messages
                except Exception as e: st.error(f"チャットエラー: {e}")

# ==========================================
# Phase 3: マッチ審査
# ==========================================
elif app_mode == "4. 書類審査":
    st.title("書類作成後: 書類審査＆推薦文作成＆面接対策")
    m_mode = st.radio("分析モード", ["1. 簡易マッチング", "2. 詳細マッチング"], horizontal=True)
    
    if m_mode == "1. 簡易マッチング":
        col1, col2 = st.columns(2)
        with col1:
            m_age = st.number_input("年齢", 18, 85, 25, key="m_age_3")
            m_ind = st.text_input("応募業種", key="m_ind_3")
        with col2:
            m_job = st.text_input("応募職種", key="m_job_3")
        
        if st.button("簡易マッチ分析を実行"):
            if not m_ind or not m_job:
                st.warning("業種と職種を入力してください。")
            else:
                prompt = f"""
あなたは採用のプロです。以下の条件から採用マッチ度（％）と、その理由を客観的に出力してください。
条件：年齢{m_age}歳、応募業種：{m_ind}、応募職種：{m_job}
出力フォーマット：
【マッチ度】
(0-100の数字)
【理由】
(簡潔に)
"""
                try:
                    resp = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    st.markdown(f"<div class='cyber-panel'>{resp.text}</div>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"APIエラー: {e}")
                    
    else:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("🏢 企業要件")
            c_url_3 = st.text_input("🔗 求人URL", key="c_url_3")
            c_info = st.text_area("求人内容", height=130)
            c_files = st.file_uploader("資料", accept_multiple_files=True, key="c_up_3")
        with col2:
            st.subheader("📄 完成書類")
            s_info = st.text_area("追加補足", height=200)
            s_files = st.file_uploader("完成書類", accept_multiple_files=True, key="s_up_3")

        if st.button("詳細審査 & 推薦文作成", type="primary"):
            if not agent_name: st.error("アドバイザー名を入力してください。")
            else:
                with st.spinner("審査中..."):
                    c_data = read_files(c_files) + "\n" + (get_url_text(c_url_3) if c_url_3 else "")
                    s_data = read_files(s_files)
                    
                    # ★修正：AIがフォーマットを崩さないよう、見出しのルールを明確化
                    prompt = f"""
あなたは凄腕ヘッドハンター兼採用担当者です。
企業要件と求職者の書類を照らし合わせ、マッチ度を％で算出し、推薦メールを作成してください。

企業情報：{c_info}\n{c_data}
求職者書類：{s_info}\n{s_data}

---
【絶対ルール】
出力は必ず以下の4つの【】見出しを含め、指定の構成で出力してください。他の【】見出しは作らないでください。

【マッチ度】
(0〜100の数字のみ)

【書類修正アドバイス】
(さらに通過率を上げるための具体的な修正点)

【面接対策】
(想定質問と回答の方向性)

【推薦文】
(以下の構成で作成)
(企業名) 採用ご担当者様
お世話になっております。キャリアアドバイザーの株式会社ライフアップの{my_name}です。
この度、○○様を推薦させていただきたく、ご連絡申し上げました。

■推薦理由
・(応募企業に活かせる強み)
・(貢献できる理由)
・(懸念点払拭があれば)
(人柄や熱意も含めて200-300字程度、AI記号「」などは禁止)

ぜひ一度、面接にてご本人とお話しいただけますと幸いです。
何卒ご検討のほど、よろしくお願い申し上げます。
"""
                    try:
                        resp = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                        res_m = resp.text
                        match_score_raw = get_section('マッチ度', res_m)
                        ms = int(re.search(r'\d+', match_score_raw).group()) if re.search(r'\d+', match_score_raw) else 0
                        st.metric("最終マッチ度", f"{ms} %")
                        st.markdown(f"#### ✍️ アドバイス\n<div class='fb-box'>{get_section('書類修正アドバイス', res_m)}</div>", unsafe_allow_html=True)
                        if ms >= 80:
                            st.success("🔥 合格ライン突破！")
                            st.code(get_section('推薦文', res_m), language="text")
                        else:
                            st.warning("⚠️ マッチ度が基準(80%)を下回っています。")
                            st.code(get_section('推薦文', res_m), language="text") # 基準を下回っていても一旦推薦文は表示する
                        st.subheader("🗣️ 面接対策")
                        st.write(get_section('面接対策', res_m))
                    except Exception as e: st.error(f"エラー: {e}")














































